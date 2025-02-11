import os
import PyPDF2
from docx import Document
from openai import OpenAI

client = OpenAI(
    api_key=os.environ.get("OPENAI_API_KEY"),
)

def translate_text(text):
    response = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"Translate the following text into English:\n\n{text}",
            }
        ],
        model="gpt-3.5-turbo",
        max_tokens=4000
    )
    return response.choices[0].message.content

def translate_pdf_to_word(pdf_path, word_path):
    pdf_reader = PyPDF2.PdfReader(pdf_path)

    document = Document()

    for page_number, page in enumerate(pdf_reader.pages):
        text = page.extract_text()
        if text:
            translated_text = translate_text(text)
            document.add_paragraph(f"Page {page_number + 1}")
            document.add_paragraph(translated_text)
            document.add_page_break()

    document.save(word_path)
    print(f'Translation saved to {word_path}')

if __name__ == "__main__":
    pdf_path = r'C:\Users\marie\Downloads\translate\translate.pdf'
    word_path = r'C:\Users\marie\Downloads\translate\translate.docx'
    translate_pdf_to_word(pdf_path, word_path)